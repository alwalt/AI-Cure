import streamlit as st
import json
from streamlit_chat import message
from util.open_chatbot import DocChatBot
from util.generate_vectors import VectorGenerator, random_token
from util.analyze_files import ImageAnalyzer, PDFAnalyzer, DocXAnalyzer
from util.gui_funcs import image_container, get_next_step_button, conversational_chat
import weaviate

chat_model = "llama3"
chat_prompt = 'chat/image_chat.txt'
embedding_model = 'sentence-transformers/all-MiniLM-L6-v2'
image_model = "llava"



if 'jsons' not in st.session_state:
	st.session_state.jsons = {}

if 'images' not in st.session_state:
	st.session_state.images = {}
	st.session_state.image_order = []

if 'pdfs' not in st.session_state:
	st.session_state.pdfs = {}
	st.session_state.pdf_order = []

if 'docxs' not in st.session_state:
	st.session_state.docxs = {}
	st.session_state.docx_order = []

if 'make_jsons' not in st.session_state:
	st.session_state.make_jsons = False
	st.session_state.make_vector_db = False
	st.session_state.enter_chat = False

if 'button_stage' not in st.session_state:
	st.session_state.button_stage = 'Generate JSON(s)'

if 'user_token' not in st.session_state:
	st.session_state.user_token = random_token()
	st.session_state.weaviate_client = weaviate.connect_to_local()
	st.session_state.vectorstore = []


# Set the title for the Streamlit app
st.title("🧬 NASA OSDR Image & PDF Analyzer")

# Create a file uploader in the sidebar
uploaded_files = st.sidebar.file_uploader("Upload File", type=['png', 'jpg', 'pdf', 'docx'], accept_multiple_files=True)
"""from langchain_core.documents import Document

if len(uploaded_files) > 0:
	import docx
	import io
	for uploaded_file in uploaded_files:
		file_type = uploaded_file.name.split('.')[-1].upper()
		if file_type == 'DOCX':
			doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
			data = [section.text for section in doc.paragraphs]
			print(data)
			data = Document(page_content='\n'.join(data), metadata={"source": uploaded_file.name},)

"""
if st.session_state.make_jsons == True:
	image_analyzer = ImageAnalyzer(image_model)
	image_analyzer.change_prompt('images/simple_json_output_prompt.txt', prompt_type='file')
	#image_list = [st.session_state.images[image_name] for image_name in st.session_state.image_order]
	for image_name in st.session_state.image_order:
		st.session_state.jsons[image_name] = {}
		while len(st.session_state.jsons[image_name]) == 0:
			output = image_analyzer.analyze_image([st.session_state.images[image_name]['img_base64']])
			if 'Error' not in output:
				st.session_state.jsons[image_name] = output
	st.session_state.make_jsons = False
elif st.session_state.make_vector_db == True:
	pdf_data = []
	for pdf in st.session_state.pdf_order:
		pdf_analyzer = PDFAnalyzer(st.session_state.pdfs[pdf])
		pdf_data+=pdf_analyzer.handlefileandingest()
	docx_data = []
	for docx in st.session_state.docx_order:
		docx_analyzer = DocXAnalyzer(st.session_state.docxs[docx])
		docx_data+=docx_analyzer.handlefileandingest()
	#
	vector_generator = VectorGenerator(embedding_model)
	image_data = vector_generator.create_documents_from_images(st.session_state.jsons) if len(st.session_state.jsons)>0 else []
	data = pdf_data+image_data+docx_data
	st.session_state.vectorstore = vector_generator.create_vector_store_weaviate(data, st.session_state.weaviate_client, st.session_state.user_token)
	st.session_state.make_vector_db = False

image_container(uploaded_files)

if st.session_state.enter_chat == True:
	chatbot = DocChatBot(chat_model, chat_prompt, embedding_model)
	chatbot.create_chat_chain(chat_prompt, st.session_state.vectorstore)

	# Initialize chat history
	if 'history' not in st.session_state:
		st.session_state['history'] = []

	# Initialize messages
	if 'generated' not in st.session_state:
		st.session_state['generated'] = ["Hello! Feel free to ask me about the pdfs & images."]

	if 'past' not in st.session_state:
		st.session_state['past'] = ["Uploaded " + ', '.join(st.session_state.image_order)]

	# Create containers for chat history and user input
	response_container = st.container()
	container = st.container()

	# User input form
	with container:
		with st.form(key='my_form', clear_on_submit=True):
			user_input = st.text_input("Query:", placeholder="Ask your question about the Images & PDFs", key='input')
			submit_button = st.form_submit_button(label='Send')

		if submit_button and user_input:
			output = conversational_chat(user_input, chatbot.chain)
			st.session_state['past'].append(user_input)
			st.session_state['generated'].append(output)

	# Display chat history
	if st.session_state['generated']:
		with response_container:
			for i in range(len(st.session_state['generated'])):
				#message(st.session_state["past"][i], is_user=True, key=str(i) + '_user', avatar_style="thumbs")
				#message(st.session_state["generated"][i], key=str(i), avatar_style="bottts")
				message(st.session_state["past"][i], is_user=True, key=str(i) + '_user', avatar_style="thumbs")
				message(st.session_state["generated"][i], key=str(i), logo="https://visualization.genelab.nasa.gov/static/images/GeneLab_patch.png")


get_next_step_button()








