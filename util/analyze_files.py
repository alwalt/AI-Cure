import ollama
import json
from langchain.document_loaders import PyMuPDFLoader
import tempfile
import docx
import io
from langchain_core.documents import Document

#from pydantic import BaseModel
#from typing import List


class ImageAnalyzer:
	def __init__(self, image_model):
		self.image_model = image_model
		self.prompt = ""
	#
	def change_model(self, new_model, model_type):
		self.image_model = new_model
	#
	def change_prompt(self, new_prompt, prompt_type='file'):
		if prompt_type == 'file':
			with open('prompts/'+new_prompt, 'r') as prompt_file:
				self.prompt = ' '.join(prompt_file.readlines())
		else:
			self.prompt = new_prompt
	#
	def change_export_type(self, new_export_type):
		self.export_type = new_export_type
	#
	def analyze_image(self, image_list):
		res = ollama.chat(
			model=self.image_model,
			messages=[
				{
					'role': 'user',
					'content': self.prompt,
					'images': image_list
				}
			]
		)
		json_output = {}
		#
		# replace usual output error
		res_text = res['message']['content']
		res_text = res_text.replace("```json", "").replace("```", "")
		try:
			json_output = json.loads(res_text)
		except json.decoder.JSONDecodeError:
			json_output['Error'] = res_text
		print(json_output)
		return json_output



class PDFAnalyzer:
	def __init__(self, uploaded_file):
		self.uploaded_file = uploaded_file

	def handlefileandingest(self):
		def format_docs(docs):
			return "\n\n".join(doc.page_content for doc in docs)

		with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
			tmp_file.write(self.uploaded_file.getvalue())
			tmp_file_path = tmp_file.name

		loader = PyMuPDFLoader(file_path=tmp_file_path)
		data = loader.load()
		return data


class DocXAnalyzer:
	def __init__(self, uploaded_file):
		self.uploaded_file = uploaded_file
	#
	def handlefileandingest(self):
		doc = docx.Document(io.BytesIO(self.uploaded_file.getvalue()))
		data = [section.text for section in doc.paragraphs]
		data = Document(page_content='\n'.join(data), metadata={"source": self.uploaded_file.name},)
		return [data]
	